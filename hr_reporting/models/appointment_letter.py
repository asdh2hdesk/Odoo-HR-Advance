# -*- coding: utf-8 -*-
import base64
import io
import logging

from odoo import api, fields, models

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_logger = logging.getLogger(__name__)


SIGNATORY_NAME = "Syed Sudam Hussain Gillani"
SIGNATORY_TITLE = "Manager- HR & EHSS"
COMPANY_LETTERHEAD = "For German Green Steel and Power Limited."

PROBATION_PERIOD = "three months"
SHORT_NOTICE_DURING_PROBATION = "ninety (90) days"
NOTICE_AFTER_CONFIRMATION = "one month"
RETIREMENT_AGE = "60 years"


class HrContractAppointmentLetter(models.Model):
    _inherit = "hr.contract"

    appointment_letter_ref = fields.Char(
        string="Appointment Letter Ref",
        copy=False,
        readonly=True,
    )
    appointment_letter_date = fields.Date(
        string="Appointment Letter Date",
        copy=False,
        default=fields.Date.context_today,
    )
    application_date = fields.Date(
        string="Application Date",
        help="Date the candidate applied; used in the appointment letter opening line.",
    )
    appointment_letter_attachment_id = fields.Many2one(
        "ir.attachment",
        string="Appointment Letter",
        copy=False,
        readonly=True,
    )

    @api.model_create_multi
    def create(self, vals_list):
        contracts = super().create(vals_list)
        for contract in contracts:
            try:
                contract._generate_appointment_letter_docx()
            except Exception:
                _logger.exception(
                    "Failed to auto-generate appointment letter for contract id=%s",
                    contract.id,
                )
        return contracts

    def action_generate_appointment_letter(self):
        self.ensure_one()
        attachment = self._generate_appointment_letter_docx()
        if not attachment:
            return False
        return {
            "type": "ir.actions.act_url",
            "url": "/web/content/%s?download=true" % attachment.id,
            "target": "self",
        }

    # ------------------------------------------------------------------
    # DOCX builder
    # ------------------------------------------------------------------
    def _generate_appointment_letter_docx(self):
        self.ensure_one()
        if not self.employee_id:
            return False

        if not self.appointment_letter_ref:
            self.appointment_letter_ref = self.env["ir.sequence"].next_by_code(
                "hr.appointment.letter"
            ) or ""
        if not self.appointment_letter_date:
            self.appointment_letter_date = fields.Date.context_today(self)
        if not self.application_date and self.create_date:
            self.application_date = self.create_date.date()

        doc = Document()
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

        self._appt_write_header(doc)
        self._appt_write_employee_block(doc)
        self._appt_write_subject_and_intro(doc)
        self._appt_write_terms(doc)
        self._appt_write_closing(doc)
        self._appt_write_annexure(doc)

        buffer = io.BytesIO()
        doc.save(buffer)

        filename = "Appointment Letter - %s%s.docx" % (
            self.employee_id.name or "Employee",
            " (%s)" % self.employee_id.employee_code if self.employee_id.employee_code else "",
        )

        attachment_vals = {
            "name": filename,
            "type": "binary",
            "datas": base64.b64encode(buffer.getvalue()),
            "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "res_model": self._name,
            "res_id": self.id,
        }
        if self.appointment_letter_attachment_id:
            self.appointment_letter_attachment_id.write(attachment_vals)
            attachment = self.appointment_letter_attachment_id
        else:
            attachment = self.env["ir.attachment"].create(attachment_vals)
            self.appointment_letter_attachment_id = attachment.id
        return attachment

    # ------------------------------------------------------------------
    # Sections
    # ------------------------------------------------------------------
    def _appt_write_header(self, doc):
        ref = self.appointment_letter_ref or ""
        date_str = self._appt_format_date(self.appointment_letter_date)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        left = table.rows[0].cells[0].paragraphs[0]
        right = table.rows[0].cells[1].paragraphs[0]
        right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        left.add_run(ref).bold = True
        run = right.add_run("DATE: %s" % date_str)
        run.bold = True

    def _appt_write_employee_block(self, doc):
        emp = self.employee_id
        title = "Mr." if (emp.gender or "") == "male" else (
            "Ms." if (emp.gender or "") == "female" else "Mr./Ms."
        )
        p = doc.add_paragraph()
        p.add_run("\n%s %s" % (title, (emp.name or "____________________").upper())).bold = True

        address_parts = [
            getattr(emp, "private_street", False),
            getattr(emp, "private_street2", False),
            getattr(emp, "private_city", False),
            emp.private_state_id.name if emp.private_state_id else False,
            getattr(emp, "private_zip", False),
        ]
        address_line = ", ".join([p for p in address_parts if p]) or "____________________"
        doc.add_paragraph("Address: %s" % address_line)

        phone = (
            getattr(emp, "private_phone", False)
            or getattr(emp, "mobile_phone", False)
            or "____________________"
        )
        doc.add_paragraph("Contact: %s" % phone)

        email = (
            getattr(emp, "private_email", False)
            or getattr(emp, "work_email", False)
            or "____________________"
        )
        doc.add_paragraph("Email: %s" % email)

    def _appt_write_subject_and_intro(self, doc):
        designation = (self.job_id.name if self.job_id else False) or "____________________"
        application_date_str = self._appt_format_date(self.application_date)
        effective_date_str = self._appt_format_date(self.date_start)

        # Centered, bold subject line
        subject = doc.add_paragraph()
        subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subject.add_run("Sub: Letter of Appointment")
        run.bold = True
        run.underline = True

        # Intro with bold inline values
        self._appt_add_runs(doc.add_paragraph(), [
            ("With reference to your application dated ", False),
            (application_date_str, True),
            (" and the subsequent interview you had with us, we are pleased to appoint you as ", False),
            (designation, True),
            (" at our Company, with effect from ", False),
            (effective_date_str, True),
            (" on the following terms and conditions:", False),
        ])

    def _appt_write_terms(self, doc):
        designation = (self.job_id.name if self.job_id else False) or "____________________"
        annual_ctc_words = self._appt_ctc_in_lakhs(getattr(self, "final_yearly_costs", 0.0) or 0.0)

        items = [
            # 1. Probation
            [
                ("Your appointment as ", False),
                (designation, True),
                (" will be initially on Probationary Basis for a period of ", False),
                (PROBATION_PERIOD, True),
                (". During the period of probation, your services may be terminated without "
                 "any notice, notice pay and without any reason being assigned. The period "
                 "of probation may be extended at the sole discretion of the Management if "
                 "considered necessary.", False),
            ],
            # 2. Gross remuneration / Annexure
            [
                ("As discussed, your gross remuneration is INR ", False),
                (annual_ctc_words, True),
                (" per annum which include direct and indirect benefit attached here your "
                 "position and a detail break up of your remuneration packaged has been "
                 "outline in ", False),
                ("Annexture-1", True),
                (".", False),
            ],
            # 3. Gratuity
            [
                ("Your retirement benefits encompass ", False),
                ("Gratuity", True),
                (" as per the policy of the Company.", False),
            ],
            # 4. Probation discontinue
            [
                ("During the probation period, in case for any reasons you wish to discontinue "
                 "your services, you may do so by giving ", False),
                (SHORT_NOTICE_DURING_PROBATION, True),
                (" notice or ", False),
                (SHORT_NOTICE_DURING_PROBATION, True),
                (" pay in lieu thereof.", False),
            ],
            # 5. Confirmation
            [
                ("On satisfactory completion of the probationary period, of which the "
                 "Management will be the sole judge, you will be confirmed in the company's "
                 "services. Unless the letter of confirmation is issued to you, you will be "
                 "deemed to be on probation.", False),
            ],
            # 6. After confirmation
            [
                ("After confirmation, the management may, at any time, terminate your services "
                 "by giving you ", False),
                (NOTICE_AFTER_CONFIRMATION, True),
                (" notice or ", False),
                (NOTICE_AFTER_CONFIRMATION, True),
                (" pay in lieu thereof. You may also resign from the services of the company "
                 "by giving ", False),
                (NOTICE_AFTER_CONFIRMATION, True),
                (" notice in writing and working actually for ", False),
                (NOTICE_AFTER_CONFIRMATION, True),
                (". However, the management can relieve you at its own discretion prior to "
                 "completion of ", False),
                (NOTICE_AFTER_CONFIRMATION, True),
                (".", False),
            ],
            # 7. Governance
            [
                ("You will be governed by the Company's rules, regulations and / or standing "
                 "orders in force from time to time with regard to the discipline and other "
                 "matters in connection with your employment with us and you will abide by "
                 "them in letter and spirit.", False),
            ],
            # 8. No other employment
            [
                ("During your employment in our establishment, you will ", False),
                ("not engage yourself in any other employment, vocation or trade", True),
                (" of any kind. If, at any time, it is found that you are violating this rule, "
                 "you will subject yourself to severe disciplinary action.", False),
            ],
            # 9. Transfer
            [
                ("Your services are liable to be transferred on full time or part time basis "
                 "to any of the existing or future office / branches of this organization or "
                 "any of its existing sister concern or the concerns newly started in future "
                 "in any part of India or shifted from one job or department to another. This "
                 "will not cause any reduction in your total emoluments or remuneration.", False),
            ],
            # 10. Overseas bond
            [
                ("In case an employee / workman is deputed by the Company on any overseas "
                 "business visit, the employee, before proceeding on such overseas visit "
                 "shall execute a ", False),
                ("service bond of three years", True),
                (" of service with the Company. In the event of the employee wanting to "
                 "separate from the Company within these three years, he / she shall "
                 "reimburse to the Company, expenses incurred on the overseas visit, on "
                 "pro-rata basis.", False),
            ],
            # 11. Fitness
            [
                ("Your continuation in the Company's employment will be subject to your being "
                 "physically and mentally fit to carry out your duties. If the company becomes "
                 "aware of the fact that you are suffering from any incurable or contagious "
                 "disease, the company may take any action as deemed fit.", False),
            ],
            # 12. False declaration
            [
                ("You will be liable to be ", False),
                ("summarily dismissed", True),
                (" if it is found that you have given a false declaration regarding name, "
                 "age, father's name, qualifications, previous service or other matters at "
                 "the time of employment.", False),
            ],
            # 13. Leave types
            [
                ("You are entitled to avail ", False),
                ("Privilege, Emergency and Medical leave", True),
                (" as per the Company norms, a copy of which is available from Human "
                 "Resources Department.", False),
            ],
            # 14. Calendar year
            [
                ("The leaves will be allotted as per the company norms in the ", False),
                ("Colander year (Jan to Dec.)", True),
                ("", False),
            ],
            # 15. Confidentiality
            [
                ("The success of our organization lies in its key's areas of performance like "
                 "process, methods, operational strategy, information, management know-how, "
                 "knowledge, etc. pertaining to manufacturing, procuring, marketing, trading "
                 "and whole gamut of is business activity. It is therefore expected, that you "
                 "will maintain ", False),
                ("utmost confidentiality", True),
                (" and under no circumstances shall directly or indirectly impart/share to any "
                 "person, any information regarding company's business, plan, progress, "
                 "prospects or affairs and shall not pledge the Credit of the Company for any "
                 "reason or purpose whatsoever.", False),
            ],
            # 16. Retirement
            [
                ("You will retire from the Company's services on the last date of the month in "
                 "which you attain the age of ", False),
                (RETIREMENT_AGE, True),
                (". In this context, please furnish us with an attested copy of your ", False),
                ("pan card and Aadhar card", True),
                (".", False),
            ],
            # 17. Address correspondence
            [
                ("All future correspondence will be sent of you at the address above in this "
                 "letter. In case of change in address, you will intimate us in writing within ",
                 False),
                ("7 days", True),
                (" from such changes.", False),
            ],
            # 18. Disputes / Jurisdiction
            [
                ("Disputes/Differences if any arising out of employment will be subject to the "
                 "Courts/Tribunal of ", False),
                ("Ahmedabad Jurisdiction", True),
                (" and shall be dealt/settled at Ahmedabad.", False),
            ],
            # 19. Reservation
            [
                ("The company reserves its right to add or deletes any of the conditions "
                 "mentioned in this letter of appointment during the continuance of your "
                 "employment without assigning any reason whatsoever.", False),
            ],
        ]

        for idx, parts in enumerate(items, start=1):
            p = doc.add_paragraph()
            num_run = p.add_run("%d. " % idx)
            num_run.bold = True
            self._appt_add_runs(p, parts)

        # Note (not numbered)
        p = doc.add_paragraph()
        note_run = p.add_run(
            "Note: - Sign in this appointment letter with the above terms and conditions "
            "would mean that you have accepted this appointment in our company unconditionally. "
            "A copy of the appointment letter, duly signed by you, shall be maintained at the "
            "Human Resources Department for the purpose of reference and records."
        )
        note_run.italic = True
        note_run.bold = True

        doc.add_paragraph("We wish you all success and happiness in your job with us.")

    def _appt_write_closing(self, doc):
        doc.add_paragraph(COMPANY_LETTERHEAD).runs[0].bold = True
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph(SIGNATORY_NAME).runs[0].bold = True
        doc.add_paragraph(SIGNATORY_TITLE)

        doc.add_paragraph()
        doc.add_paragraph("I hereby accept all the above terms & conditions.")
        sig_line = doc.add_paragraph()
        sig_line.add_run("Signature of Employee\t\t\t\t\t\t\tDate:")
        doc.add_paragraph(
            "(Mr. / Mrs. / Ms. _______________________)"
        )

    def _appt_write_annexure(self, doc):
        doc.add_page_break()

        # Pull amounts from contract_salary_config lines if available
        components = [
            ("Basic salary", "BASIC"),
            ("House rent allowances", "HRA"),
            ("Conveyance allowances", "CONVEYANCE"),
            ("Medical allowances", "MEDICAL"),
            ("Incentive allowances", "INCENTIVE"),
            ("Other allowances", "OTHER"),
        ]
        deductions = [
            ("Pt", "PT"),
            ("Bonus", "BONUS"),
            ("IT", "IT"),
        ]
        comp_rows = [(label, self._appt_amount(code)) for label, code in components]
        ded_rows = [(label, self._appt_amount(code)) for label, code in deductions]

        gross_salary = (
            getattr(self, "gross_salary", False)
            or getattr(self, "monthly_yearly_costs", False)
            or sum(amount for _label, amount in comp_rows)
        )
        annual_ctc = getattr(self, "final_yearly_costs", 0.0) or (gross_salary * 12)
        monthly_ctc = getattr(self, "monthly_yearly_costs", 0.0) or gross_salary
        total_deduction = sum(amount for _label, amount in ded_rows)
        total_payable = (
            getattr(self, "inhand_salary", False)
            or (gross_salary - total_deduction)
        )

        # Build a 4-col table: components | amount | deductions/totals | amount
        max_rows = max(len(comp_rows), len(ded_rows))
        # header(1) + name/CTC(1) + designation/PM CTC(1) + column-headers(1)
        # + body(max_rows) + spacer(1) + Gross/Total deduction(1) + Total payable(1)
        # + note(1)
        total_rows = 4 + max_rows + 1 + 2 + 1
        table = doc.add_table(rows=total_rows, cols=4)
        table.style = "Table Grid"

        # Header row (merged)
        header_cell = table.cell(0, 0).merge(table.cell(0, 1)).merge(
            table.cell(0, 2)
        ).merge(table.cell(0, 3))
        header_p = header_cell.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_p.add_run("ANEXTURE-1")
        run.bold = True
        run.font.size = Pt(13)

        # Row 1: NAME | <name> | ANNUAL CTC | <ctc>
        self._appt_set_cell(table.cell(1, 0), "NAME", bold=True)
        self._appt_set_cell(table.cell(1, 1), (self.employee_id.name or "").upper())
        self._appt_set_cell(table.cell(1, 2), "ANNUAL CTC", bold=True)
        self._appt_set_cell(table.cell(1, 3), self._appt_money(annual_ctc))

        # Row 2: DESIGNATION | <job> | P.M CTC | <monthly>
        self._appt_set_cell(table.cell(2, 0), "DESIGNATION", bold=True)
        self._appt_set_cell(table.cell(2, 1), self.job_id.name or "")
        self._appt_set_cell(table.cell(2, 2), "P.M CTC", bold=True)
        self._appt_set_cell(table.cell(2, 3), self._appt_money(monthly_ctc))

        # Row 3: column headers
        self._appt_set_cell(table.cell(3, 0), "Components", bold=True)
        self._appt_set_cell(table.cell(3, 1), "Fix. Amount", bold=True)
        self._appt_set_cell(table.cell(3, 2), "Heads", bold=True)
        self._appt_set_cell(table.cell(3, 3), "Deductions", bold=True)

        body_start = 4
        for i in range(max_rows):
            row_idx = body_start + i
            if i < len(comp_rows):
                label, amount = comp_rows[i]
                self._appt_set_cell(table.cell(row_idx, 0), label)
                self._appt_set_cell(table.cell(row_idx, 1), self._appt_money(amount))
            if i < len(ded_rows):
                label, amount = ded_rows[i]
                self._appt_set_cell(table.cell(row_idx, 2), label)
                self._appt_set_cell(
                    table.cell(row_idx, 3),
                    self._appt_money(amount) if amount else "",
                )

        spacer_idx = body_start + max_rows
        gross_idx = spacer_idx + 1
        payable_idx = spacer_idx + 2
        note_idx = payable_idx + 1

        # Gross row
        self._appt_set_cell(table.cell(gross_idx, 0), "Gross salary", bold=True)
        self._appt_set_cell(table.cell(gross_idx, 1), self._appt_money(gross_salary), bold=True)
        self._appt_set_cell(table.cell(gross_idx, 2), "Total deduction", bold=True)
        self._appt_set_cell(table.cell(gross_idx, 3), self._appt_money(total_deduction), bold=True)

        # Total payable row
        self._appt_set_cell(table.cell(payable_idx, 2), "Total payable", bold=True)
        self._appt_set_cell(table.cell(payable_idx, 3), self._appt_money(total_payable), bold=True)

        # Note row (merged)
        note_cell = table.cell(note_idx, 0).merge(table.cell(note_idx, 1)).merge(
            table.cell(note_idx, 2)
        ).merge(table.cell(note_idx, 3))
        note_p = note_cell.paragraphs[0]
        note_run = note_p.add_run(
            "NOTE: - Bonus depends on your monthly attendance, it's not a fix amount"
        )
        note_run.italic = True

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------
    def _appt_amount(self, code):
        getter = getattr(self, "get_salary_breakdown_amount", None)
        if getter:
            try:
                return float(getter(code) or 0.0)
            except Exception:
                return 0.0
        # Fallback: scan salary_structure_line_ids if present
        lines = getattr(self, "salary_structure_line_ids", None)
        if lines:
            line = lines.filtered(lambda l: (l.code or "") == code)[:1]
            return float(line.amount_monthly) if line else 0.0
        return 0.0

    @staticmethod
    def _appt_money(value):
        try:
            return "{:,.0f}".format(float(value or 0.0))
        except (TypeError, ValueError):
            return ""

    @staticmethod
    def _appt_format_date(value):
        if not value:
            return "____________________"
        return value.strftime("%d/%m/%Y")

    @staticmethod
    def _appt_ctc_in_lakhs(amount):
        try:
            amount = float(amount or 0.0)
        except (TypeError, ValueError):
            amount = 0.0
        if amount <= 0:
            return "____________________"
        lakhs = amount / 100000.0
        return "₹%.2f Lakhs" % lakhs

    @staticmethod
    def _appt_add_runs(paragraph, parts):
        """Append a sequence of (text, is_bold) runs to a paragraph."""
        for text, is_bold in parts:
            if not text:
                continue
            run = paragraph.add_run(text)
            run.bold = bool(is_bold)

    @staticmethod
    def _appt_set_cell(cell, text, bold=False):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.text = ""
        run = p.add_run(text or "")
        run.bold = bold
        run.font.size = Pt(11)
