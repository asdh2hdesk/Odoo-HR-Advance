from odoo import models
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
import base64
from datetime import date
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

class HrCustomFormEleven(models.Model):
    _inherit = "hr.custom.form.form11"
    
    def action_generate_excel_report(self):
        self.ensure_one()

        doc = Document()
        
        # ================= COMMON FONT =================
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # ============ Page 1 ===========
        # TABLE
        t = doc.add_table(rows=1, cols=9)
        t.autofit = False
        t.alignment = WD_TABLE_ALIGNMENT.CENTER 

        cell = t.cell(0, 0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        p = cell.paragraphs[0]
        run = p.add_run()
        # Get the absolute path to the image file
        module_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        image_path = os.path.join(module_path, "static", "images", "emp.png")
        if os.path.exists(image_path):
            run.add_picture(image_path, width=Cm(1.6))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ---------- CENTER TEXT ----------
        t.cell(0,1).merge(t.cell(0,7))
        cell = t.cell(0, 1)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p.text = (
            "EMPLOYEES' PROVIDENT FUND ORGANISATION\n"
            "(Ministry of Labour & Employment, Govt. of India)\n"
            "Head Office\n"
            "Bhavishya Nidhi Bhawan, 14-Bhikaiji Cama Place, New Delhi-110066\n"
            "www.epfindia.gov.in    www.epfindia.nic.in\n"
            "Telephone: 011-26713254  Fax: 011-26166609\n"
            "Email: acc.fa.imc@epfindia.gov.in"
        )

        # ---------- RIGHT IMAGE ----------
        cell = t.cell(0, 8)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        p = cell.paragraphs[0]
        run = p.add_run()
        # Get the absolute path to the image file
        module_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        image_path = os.path.join(module_path, "static", "images", "pandit.png")
        if os.path.exists(image_path):
            run.add_picture(image_path, width=Cm(1.6))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER        
        
        doc.add_paragraph(f"Date:{date.today().strftime('%d-%m-%Y') or ''}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph("No: Manual/Amendment/2011\n")
        
        doc.add_paragraph("To\n")
   
        doc.add_paragraph(
            "       All Addi. CPFC (HQ/Zone),\n"
            "       Regional P.F. Commissioners-incharge of\n"
            "       Regional Offices.\n\n"
            "       Subject: Introduction of Composite Declaration Form (F-11)\n"
            "Sir,\n"
        )
        
        doc.add_paragraph(
            "       The Central Provident Fund Commissioner by exercising the powers conferred under para"
            "36(7) read alongwith the provisions of para 34 and 57 of EPF Scheme, 1952 and para 24 of"
            "Employees' Pension Scheme, 1995 has ordered the introduction of Composite Declaration Form"
            "(F-11) by replacing the existing New Form-11 and the same is enclosed as Annexure.\n" 
        )
        
        doc.add_paragraph("Yours faithfully,\n").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph("Encl: As above\n").alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph("(Udita Chowdhary)\nAddi. Central P.F. Commissioner (F&A)\n\n\n\n\n\n").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        
        # ================= Page 2 =================.
        # TABLE
        t = doc.add_table(rows=1, cols=9)
        t.autofit = False
        t.alignment = WD_TABLE_ALIGNMENT.CENTER 

        cell = t.cell(0, 0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        p = cell.paragraphs[0]
        run = p.add_run()
        module_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        image_path = os.path.join(module_path, "static", "images", "emp.png")
        if os.path.exists(image_path):
            run.add_picture(image_path, width=Cm(1.6))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


        # ---------- CENTER TEXT ----------
        t.cell(0,1).merge(t.cell(0,7))
        cell = t.cell(0, 1)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p.text = (
            "Employees' Provident Fund Organization\n"
            "(Minisoy of Labour & Employment, Govt. Of India)\n"
            "Head Office\n"
            "Bhavishya Nidhi Bhawan, 14-Bhikaiji Cama Place, New Delhi-110066\n"
            "www.epfindia.gov.in             www.epfindia.nic.in\n"
            "Telephone: 011- 26713254 Fax: OJ l-26166609 Email: acc.fa.imc(a:epfindia.gov.in\n"
        )

        # ---------- RIGHT IMAGE ----------
        cell = t.cell(0, 8)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        p = cell.paragraphs[0]
        run = p.add_run()
        module_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        image_path = os.path.join(module_path, "static", "images", "pandit.png")
        if os.path.exists(image_path):
            run.add_picture(image_path, width=Cm(1.6))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        

        doc.add_paragraph(f"Date:{date.today().strftime('%d-%m-%Y') or ''}").alignment = WD_ALIGN_PARAGRAPH.RIGHT  
        
        doc.add_paragraph("No: Manual/Amendment/2011\n")
        

        doc.add_paragraph(
            "ORDER\n"
            "Introduction of New Form 11\n"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER 


        doc.add_paragraph(
            "           The Employees' Provident Fund Organization has embarked upon next phase of\n"
            "e-governance reforms with a view to make its services available to its stakeholders. EPFO\n"
            "has recently introduced a single page Composite Claim Form (Aadhaar/Non-Aadhaar) and"
            "Composite Oaim Form for death cases by replacing multiple forms for settlement of claims.\n\n" 
            
            "2.     In exercise of powers conferred under para 36(7) read alongwith the provisions of" 
            "para 34 and 57 of EPF Scheme, 1952 and para 24 of Employees' Pension Scheme, 1995, the"
            "introduction of Composite Declaration Form (F-11) is ordered with immediate effect by replacing"
            "the existing New Form-11.\n\n" 
            
            "3.      The Composite Declaration Form will also replace Form No. 13 in all cases of auto" 
            "transfer vide order No. Manual/Amendment/2011/13326' dated 20.09.2017.\n\n"
        )

        doc.add_paragraph("(Dr. V.P. Joy)\nCentral Provident Fund Commissioner").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph("Encl: Composite Declaration Form-11\n\n\n\n\n\n\n").alignment = WD_ALIGN_PARAGRAPH.LEFT

        # ================= Page 3 =================
        doc.add_paragraph("www.epfindia.gov.in\n").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # TABLE
        t = doc.add_table(rows=1, cols=9)
        t.autofit = False
        t.alignment = WD_TABLE_ALIGNMENT.CENTER 

        cell = t.cell(0, 0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        p = cell.paragraphs[0]
        run = p.add_run()
        module_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        image_path = os.path.join(module_path, "static", "images", "emp.png")
        if os.path.exists(image_path):
            run.add_picture(image_path, width=Cm(1.6))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


        # ---------- CENTER TEXT ----------
        t.cell(0,1).merge(t.cell(0,8))
        cell = t.cell(0, 1)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p.text = (
            "Composite Declaration Form -11\n"
            "(To be retained by the employer for future reference)\n"
            "EMPLOYEES' PROVIDENT FUND ORGANISATION\n"
            "Employees' Provident Funds Scheme, 1952 (Paragraph 34 & 57) &\n"
            "Employees' Pension Scheme, 1995 (Paragraph 24)\n"
            "(Declaration by a person taking op employment in any establishment on which EPF Scheme, 1952 and /or EPS, 1995 is applicable)\n"
        )
        

        # ================= MAIN DETAILS TABLE =================
        table = doc.add_table(rows=16, cols=2)
        table.style = 'Table Grid'

        rows = [
            "Name of the Member",
            "Father's Name",
            "Spouse's Name",
            "Date of Birth (DD/MM/YYYY)",
            "Gender (Male/Female/Transgender)",
            "Marital Status (Married/Unmarried/Widow/Widower/Divorcee)",
            "Email ID",
            "Mobile No.",
            "Present employment details:\nDate of joining in the current establishment (DD/MM/YYYY)",
            "KYC Details: (attach selfattested copies of following KYCs)",
            "Bank Account No.",
            "IFS Code of the branch:",
            "AADHAR Number",
            "Permanent Account Number (PAN), if available",
            "Whether earlier a member of Employees' Provident Fund Scheme, 1952",
            "Whether earlier a member of Employees' Pension Scheme, 1995",
        ]
        
        data = [
            self.employee_id.name or "",
            self.father_name or "",
            self.spouse_name or "",
            self.date_of_birth.strftime('%d-%m-%Y')  or "",
            self.gender or "",
            self.marital_status or "",
            self.email or "",
            self.mobile or "",
            self.present_joining_date.strftime('%d-%m-%Y') or "",
            "",
            self.bank_account_no or "",
            self.bank_ifsc or "",
            self.aadhaar_number or "",
            self.pan_number or "",
            self.member_epf_before or "",
            self.member_eps_before or "",  
        ]

        for i, label in enumerate(rows):
            table.cell(i, 0).text = label
            table.rows[i].height = Inches(0.3)
        
        for i, rec in enumerate(data):
            table.cell(i, 1).text = str(rec)
        
        

        # ================= PREVIOUS EMPLOYMENT TABLE =================
        doc.add_paragraph("\n\n\n\n\n\nPrevious employment details: (if Yes to 9 AND/OR 10 above I - Un-exempted")

        prev1 = doc.add_table(rows=len(self.prev_unexempted_line_ids)+1, cols=8)
        prev1.autofit = True
        prev1.style = 'Table Grid'

        headers = [
            "Establishment Name & Address",
            "Universal Account Number",
            "PF Account Number ",
            "Date of joining (DD/MM/YYYY)",
            "Date of exit (DD/MM/YYYY ",
            "Scheme Certificate No. (if issued)",
            "PPO Number (if issued)",
            "Non Contributory Period (NCP) Days",
        ]
        
        for i, h in enumerate(headers):
            prev1.cell(0, i).text = h
            
        row_index = 1

        for rec in self.prev_unexempted_line_ids:
            prev1.cell(row_index, 0).text = rec.establishment or ""
            prev1.cell(row_index, 1).text = str(rec.uan) or ""
            prev1.cell(row_index, 2).text = str(rec.pf_account) or ""
            prev1.cell(row_index, 3).text = str(rec.joining_date.strftime('%d-%m-%Y')) or ""
            prev1.cell(row_index, 4).text = str(rec.exit_date.strftime('%d-%m-%Y')) or ""
            prev1.cell(row_index, 5).text = str(rec.scheme_certificate) or ""
            prev1.cell(row_index, 6).text = str(rec.ppo_number) or ""
            prev1.cell(row_index, 7).text = str(rec.ncp_days) or ""

            row_index += 1

        # ================= PREVIOUS EMPLOYMENT Detail =================
        doc.add_paragraph("\n\nPrevious employment details: (if Yes to 9 AND/OR 10 above)- For Exempted Trusts")

        prev2 = doc.add_table(rows=len(self.prev_exempted_line_ids)+1, cols=7)
        prev2.autofit = True
        prev2.style = 'Table Grid'

        headers = [
            "Name & Address of the Trust",
            "UAN",
            "Member EPS A/c Number",
            "Date of joining (DD/MM/YYYY)",
            "Date of exit (DD/MM/YYYY)",
            "Scheme Certificate No.(if issued)",
            "Non Contributory Period (NCP) Days",
        ]

        for i, h in enumerate(headers):
            prev2.cell(0, i).text = h

        # ===== Data Rows =====
        row_index = 1

        for rec in self.prev_exempted_line_ids:
            prev2.cell(row_index, 0).text = rec.trust_name or ""
            prev2.cell(row_index, 1).text = str(rec.uan) or ""
            prev2.cell(row_index, 2).text = str(rec.eps_account) or ""
            prev2.cell(row_index, 3).text = str(rec.joining_date.strftime('%d-%m-%Y')) or ""
            prev2.cell(row_index, 4).text = str(rec.exit_date.strftime('%d-%m-%Y')) or ""
            prev2.cell(row_index, 5).text = str(rec.scheme_certificate) or ""
            prev2.cell(row_index, 6).text = str(rec.ncp_days) or ""

            row_index += 1

    
        # ================= INTERNATIONAL WORKER =================
        doc.add_paragraph("\n")

        t1 = doc.add_table(rows=4, cols=2)
        t1.style = 'Table Grid'
        
        header = [
            "International Worker:",
            "If yes, state country of origin (India/Name of other country)",
            "Passport No:",
            "Validity of passport [(DD/MM/YYYY) to (DD/MM/YYYY)]",
        ]

        passport_validity = ""

        if self.passport_valid_from or self.passport_valid_to:
            from_date = self.passport_valid_from.strftime('%d-%m-%Y') if self.passport_valid_from else ''
            to_date = self.passport_valid_to.strftime('%d-%m-%Y') if self.passport_valid_to else ''
            passport_validity = f"{from_date} to {to_date}"
        
        data = [
            self.international_worker or '',
            self.origin_country or '',
            self.passport_no or '',
            passport_validity,
        ]

        for i, label in enumerate(header):
            t1.cell(i, 0).text = label
            t1.rows[i].height = Inches(0.3)

        for i, rec in enumerate(data):
            t1.cell(i, 1).text = str(rec)

        # ================= UNDERTAKING =================
        u = doc.add_paragraph("\nUNDERTAKING")
        u.runs[0].bold = True
        u.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            "1. Certified that the particulars are true to the best of my knowledge.\n"
            "2. I authorize EPFO to use my Aadhar for verification/authentication/e-KYC purpose for service delivery.\n"
            "3. Kindly transfer the funds and service details, if applicable, from the previous PF account as declared above to the present P.F. Account as I am an Aadhar verified employee in my previous PF Accounl * \n"
            "4. In case of changes in above details, the same will be intimated to employer at the earliest.\n"
        )
        
        doc.add_paragraph("Date:\nPlace:")
        doc.add_paragraph("Signature of Member\n").alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # ================= EMPLOYER DECLARATION =================
        d = doc.add_paragraph("\nDECLARATION BY PRESENT EMPLOYER\n")
        d.runs[0].bold = True
        d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(
            f"A.   The member Mr/Ms/Mrs {self.employee_id.name or ''} has joined on {str(self.present_joining_date.strftime('%d-%m-%Y') or '')} "
            f"and has been allotted PF No _______________ and UAN {self.employee_id.l10n_in_uan}."
        )

        doc.add_paragraph(
            "B.   In case the person was earlier not a member of EPF Scheme, 1952 and EPS, 1995:\n"
            "       Please Tick the Appropriate Option:\n"
            "       The KYC details of the above member in the UAN database\n"
            "           ☐ Have not been uploaded\n"
            "           ☐ Have been uploaded but not approved \n"
            "           ☐ Have been uploaded and approved with DSC/e-sign.\n"
        )
        
        doc.add_paragraph(
            "C.   In case the person was earlier a member of EPF Scheme, 1952 and EPS, 1995\n"
            "       Please Tick the Appropriate Option\n"
            "           ☐ The KYC details of the above member in the UAN database have been approved with\n              E-sign/Digital Signature Certificate and transfer request has been generated on portal.\n"
            "           ☐ The previous Account of the member is not Aadhar verified and hence physical transfer\n              form shall be initiated.\n\n"
            f"Date:{date.today().strftime('%d-%m-%Y') or ''}\n"
        )
        
        doc.add_paragraph("Signature of Employer with Seal of\nEstablishment\n\n").alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph("* Auto transfer of previous PF account would be possible in respect of Aadhar verified employees only."
                          "Other employees are requested to file physical claim (Form-13) for transfer of account from the"
                          "previous establishment.")


        buffer = io.BytesIO()
        doc.save(buffer)
        
        filename = (
            f"Form 11 - {self.employee_id.name or ''}({self.employee_id.employee_code or ''}).docx"
            if self.employee_id.employee_code
            else f"Form 11 -{self.employee_id.name or ''}.docx"
        )

        # ===== Create Attachment =====
        attachment = self.env['ir.attachment'].create({
            'name': filename,
            'type': 'binary',
            'datas': base64.b64encode(buffer.getvalue()),
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'res_model': self._name,
            'res_id': self.id,
        })

        # ===== Download Action =====
        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'self',
        }
