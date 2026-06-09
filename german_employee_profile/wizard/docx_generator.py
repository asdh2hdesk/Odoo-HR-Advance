# -*- coding: utf-8 -*-
import base64
import io

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HEADER_FILL = "F5CBA7"
BORDER_COLOR = "B0B0B0"


def _set_cell_border(cell, color=BORDER_COLOR, size=6, style="single"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn("w:%s" % edge))
        if element is None:
            element = OxmlElement("w:%s" % edge)
            tc_borders.append(element)
        element.set(qn("w:val"), style)
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def _clear_cell_border(cell):
    _set_cell_border(cell, color="FFFFFF", size=0, style="nil")


def _set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def _set_text(cell, text, *, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)


def _format_cell(
    cell,
    text,
    *,
    bold=False,
    size=8,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    fill=None,
    border=True,
    width=None,
):
    _set_text(cell, text, bold=bold, size=size, align=align)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        _set_cell_shading(cell, fill)
    if border:
        _set_cell_border(cell)
    else:
        _clear_cell_border(cell)
    if width:
        _set_cell_width(cell, width)


def _section_heading(table, row_index, title):
    merged = table.cell(row_index, 0)
    for col in range(1, len(table.columns)):
        merged = merged.merge(table.cell(row_index, col))
    _format_cell(
        merged,
        title,
        bold=True,
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        fill=HEADER_FILL,
    )


def _set_table_font(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def _add_spacer(doc, pts):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)


def _add_company_logo(doc, logo_b64):
    if not logo_b64:
        return

    try:
        image_bytes = base64.b64decode(logo_b64)
    except Exception:
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(image_bytes), width=Inches(2.6))
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)


def _add_image_to_cell(cell, image_b64, *, width_inches, fallback_text):
    _set_text(cell, "")
    if not image_b64:
        _set_text(cell, fallback_text, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        return

    try:
        image_bytes = base64.b64decode(image_b64)
    except Exception:
        _set_text(cell, fallback_text, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        return

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    try:
        run.add_picture(io.BytesIO(image_bytes), width=Inches(width_inches))
    except Exception:
        _set_text(cell, fallback_text, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)


def _build_top_info(doc, data):
    table = doc.add_table(rows=6, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1.4, 2.2, 1.8, 1.9]
    for row in table.rows:
        for idx, width in enumerate(widths):
            _set_cell_width(row.cells[idx], width)

    left_rows = [
        ("EMPLOYEE CODE :-", data["employee_code"]),
        ("DATE :-", data["date"]),
        ("DEPARTMENT:-", data["department"]),
        ("EMPLOEE POST OF-", data["post"]),
        ("BLOOD GROOP:-", data["blood_group"]),
    ]

    for idx, (label, value) in enumerate(left_rows):
        _format_cell(table.cell(idx, 0), label, bold=True, border=False, width=widths[0])
        _format_cell(table.cell(idx, 1), value, border=False, width=widths[1])

    photo_cell = table.cell(0, 2).merge(table.cell(5, 3))
    _format_cell(photo_cell, "", align=WD_ALIGN_PARAGRAPH.RIGHT, width=widths[2] + widths[3])
    _add_image_to_cell(photo_cell, data.get("employee_photo"), width_inches=1.10, fallback_text="PASSPORT SIZE\nPHOTO")
    return table


def _build_personal_table(doc, data):
    table = doc.add_table(rows=14, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _section_heading(table, 0, "PERSONAL DETAILS")

    widths = [0.45, 1.4, 1.7, 0.45, 1.55, 2.15]
    for row in table.rows:
        for idx, width in enumerate(widths):
            _set_cell_width(row.cells[idx], width)

    _format_cell(table.cell(1, 0), "SR NO", bold=True, fill=HEADER_FILL, align=WD_ALIGN_PARAGRAPH.CENTER)
    _format_cell(table.cell(1, 1), "PARTICULAR", bold=True, fill=HEADER_FILL, align=WD_ALIGN_PARAGRAPH.CENTER)
    merged = table.cell(1, 2).merge(table.cell(1, 5))
    _format_cell(merged, "DETAILS", bold=True, fill=HEADER_FILL, align=WD_ALIGN_PARAGRAPH.CENTER)

    merged_rows = [
        ("1", "NAME", data["name"]),
        ("2", "FATHER'S/HUSBAND\nNAME", data["father_name"]),
        ("3", "MOTHER'S NAME", data["mother_name"]),
        ("4", "ADDRESS", data["address"]),
        ("5", "CONTACT NO.", data["contact"]),
        ("6", "E MAIL ID", data["email"]),
    ]
    row_idx = 2
    for sr_no, label, value in merged_rows:
        _format_cell(table.cell(row_idx, 0), sr_no, align=WD_ALIGN_PARAGRAPH.CENTER, size=7.5)
        _format_cell(table.cell(row_idx, 1), label, size=7.5)
        merged = table.cell(row_idx, 2).merge(table.cell(row_idx, 5))
        _format_cell(merged, value, size=7.5)
        row_idx += 1

    dual_rows = [
        ("7", "LANGAUGE KNOWN", data["language"], "13", "FOOD PREFRENCE", data["food_preference"]),
        ("8", "BIRTH DATE", data["birth_date"], "14", "AADHAAR NO.", data["aadhaar"]),
        ("9", "GENDER", data["gender"], "15", "PAN CARD NO.", data["pan"]),
        ("10", "MARITAL STATUS", data["marital"], "16", "NATIONALITY", data["nationality"]),
        ("11", "HEIGHT", data["height"], "17", "WEIGHT", data["weight"]),
        ("12", "BANK ACCOUNT NO", data["bank_account"], "18", "UAN.NO", data["uan"]),
    ]
    for values in dual_rows:
        sr_l, part_l, detail_l, sr_r, part_r, detail_r = values
        _format_cell(table.cell(row_idx, 0), sr_l, align=WD_ALIGN_PARAGRAPH.CENTER, size=7.5)
        _format_cell(table.cell(row_idx, 1), part_l, size=7.5)
        _format_cell(table.cell(row_idx, 2), detail_l, size=7.5)
        _format_cell(table.cell(row_idx, 3), sr_r, align=WD_ALIGN_PARAGRAPH.CENTER, size=7.5)
        _format_cell(table.cell(row_idx, 4), part_r, size=7.5)
        _format_cell(table.cell(row_idx, 5), detail_r, size=7.5)
        row_idx += 1
    return table


def _build_simple_table(doc, title, headers, rows, widths, centered_cols=None):
    centered_cols = centered_cols or set()
    table = doc.add_table(rows=len(rows) + 2, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _section_heading(table, 0, title)
    for row in table.rows:
        for idx, width in enumerate(widths):
            _set_cell_width(row.cells[idx], width)

    for idx, header in enumerate(headers):
        _format_cell(
            table.cell(1, idx),
            header,
            bold=True,
            fill=HEADER_FILL,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=8,
        )

    for row_idx, row_values in enumerate(rows, start=2):
        for col_idx, value in enumerate(row_values):
            _format_cell(
                table.cell(row_idx, col_idx),
                value,
                size=7.5,
                align=WD_ALIGN_PARAGRAPH.CENTER if col_idx in centered_cols else WD_ALIGN_PARAGRAPH.LEFT,
            )
    return table


def _build_salary_table(doc, data):
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1.8, 1.8, 1.8, 1.86]
    for row in table.rows:
        for idx, width in enumerate(widths):
            _set_cell_width(row.cells[idx], width)

    _format_cell(table.cell(0, 0), "LAST SALARY", bold=True, fill=HEADER_FILL)
    _format_cell(table.cell(0, 1), data["last_salary"], size=7.5)
    _format_cell(table.cell(0, 2), "EXPECTED SALARY", bold=True, fill=HEADER_FILL)
    _format_cell(table.cell(0, 3), data["expected_salary"], size=7.5)

    _format_cell(table.cell(1, 0), "JOB SKILL", bold=True, fill=HEADER_FILL)
    merged = table.cell(1, 1).merge(table.cell(1, 3))
    _format_cell(merged, data["job_skill"], size=7.5)
    return table


def _build_interview_table(doc, data):
    table = doc.add_table(rows=8, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [2.1, 3.8, 1.4]
    for row in table.rows:
        for idx, width in enumerate(widths):
            _set_cell_width(row.cells[idx], width)

    merged = table.cell(0, 0).merge(table.cell(0, 1))
    _format_cell(merged, "", border=False)
    _format_cell(
        table.cell(0, 2),
        "SIGNATURE",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        border=False,
    )

    rows = [
        ("INTERVIEWER NAME", data["interviewer"]),
        ("SELECTED/REJECTED", data["selected_rejected"]),
        ("DEPARTMENT", data["dept_interview"]),
        ("DATE OF JOINING", data["date_joining"]),
        ("FINAL SALARY", data["final_salary"]),
        ("SELECTION DATE", data["selection_date"]),
        ("REFERENCE", data["reference"]),
    ]
    for idx, (label, value) in enumerate(rows, start=1):
        _format_cell(table.cell(idx, 0), label, bold=True, fill=HEADER_FILL, size=7.5)
        _format_cell(table.cell(idx, 1), value, size=7.5)
        _format_cell(table.cell(idx, 2), "", size=7.5)
    return table


def _build_signature_table(doc):
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [2.4, 2.4, 2.66]
    for row in table.rows:
        for idx, width in enumerate(widths):
            _set_cell_width(row.cells[idx], width)

    headers = ["DEPARTMENT HEAD", "GENERAL MANAGER", "H.R.MANAGER"]
    for idx, value in enumerate(headers):
        _format_cell(
            table.cell(0, idx),
            value,
            bold=True,
            size=9,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            border=False,
        )
        _format_cell(
            table.cell(1, idx),
            "SIGNATURE",
            bold=True,
            size=8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            border=False,
        )
    return table


def build_employee_profile_docx(data):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    _add_company_logo(doc, data.get("company_logo"))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(data.get("company_name") or "GERMAN GREEN STEEL AND POWER LIMITED")
    title_run.bold = True
    title_run.font.name = "Niagara Solid"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Niagara Solid")
    title_run.font.size = Pt(48)
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    rev = doc.add_paragraph()
    rev.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rev.paragraph_format.space_before = Pt(2)
    rev.paragraph_format.space_after = Pt(2)
    rev_run = rev.add_run(data.get("revision_label") or "REV/00/00-00-0000")
    rev_run.font.name = "Arial"
    rev_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    rev_run.font.size = Pt(7)

    _build_top_info(doc, data)
    _add_spacer(doc, 4)
    _build_personal_table(doc, data)
    _add_spacer(doc, 4)
    _build_simple_table(
        doc,
        "QUALIFICATION DETAILS",
        ["EXAM", "YEAR OF PASSING", "BOARD/UNI.", "PERCENTAGE"],
        [[q["exam"], q["year"], q["board"], q["percentage"]] for q in data["qualifications"]],
        [1.9, 1.9, 1.9, 1.9],
        centered_cols={1, 3},
    )
    _add_spacer(doc, 4)
    _build_simple_table(
        doc,
        "EXPERIENCE DETAILS",
        ["SR NO", "COMPANY NAME", "DESIGNATION", "TIME PERIOD", "REASON FOR"],
        [[str(idx + 1), e["company"], e["designation"], e["period"], e["reason"]] for idx, e in enumerate(data["experiences"])],
        [0.45, 2.2, 1.8, 1.45, 1.7],
        centered_cols={0, 3, 4},
    )
    _add_spacer(doc, 4)
    _build_simple_table(
        doc,
        "FAMILY CONTACT DETAILS",
        ["NAME", "CONTACT NO.", "RELATION", "ADDRESS"],
        [[c["name"], c["contact"], c["relation"], c["address"]] for c in data["family_contacts"]],
        [1.9, 1.9, 1.9, 1.9],
        centered_cols={1, 2},
    )
    _add_spacer(doc, 4)
    _build_simple_table(
        doc,
        "EMERGENCY CONTACT DETAILS",
        ["NAME", "CONTACT NO.", "RELATION", "ADDRESS"],
        [[data["emergency"]["name"], data["emergency"]["contact"], data["emergency"]["relation"], data["emergency"]["address"]]],
        [1.9, 1.9, 1.9, 1.9],
        centered_cols={1, 2},
    )
    _add_spacer(doc, 4)
    _build_salary_table(doc, data)
    _add_spacer(doc, 6)
    _build_interview_table(doc, data)
    _add_spacer(doc, 14)
    _build_signature_table(doc)

    _set_table_font(doc.tables[0])
    for table in doc.tables[1:]:
        _set_table_font(table)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
